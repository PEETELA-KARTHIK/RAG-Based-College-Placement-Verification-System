"""
Format-aware text extraction for PDF, DOCX, TXT, and XLSX files.
Each extractor returns raw text suitable for downstream cleaning and chunking.
"""

import os
from pathlib import Path

from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from openpyxl import load_workbook


def extract_pdf(file_path: str) -> str:
    """Extract text from a PDF file, page by page."""
    reader = PdfReader(file_path)
    pages = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            pages.append(text.strip())
    return "\n\n".join(pages)


def extract_docx(file_path: str) -> str:
    """Extract text from a DOCX file, paragraph by paragraph."""
    doc = DocxDocument(file_path)
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)

    return "\n\n".join(paragraphs)


def extract_txt(file_path: str) -> str:
    """Extract text from a plain text file with encoding fallback."""
    for encoding in ("utf-8", "utf-8-sig", "latin-1", "cp1252"):
        try:
            with open(file_path, "r", encoding=encoding) as f:
                return f.read()
        except (UnicodeDecodeError, UnicodeError):
            continue
    raise ValueError(f"Unable to decode text file: {file_path}")


def extract_xlsx(file_path: str) -> str:
    """
    Extract text from an Excel file in a row-wise, column-aware format.
    Each row is stringified as 'Column1: Value1 | Column2: Value2 | ...'.
    """
    wb = load_workbook(file_path, read_only=True, data_only=True)
    sheet_texts = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            continue

        # First row as headers
        headers = [str(h).strip() if h is not None else f"Col_{i}"
                   for i, h in enumerate(rows[0])]

        row_strings = [f"--- Sheet: {sheet_name} ---"]
        for row in rows[1:]:
            pairs = []
            for header, value in zip(headers, row):
                if value is not None and str(value).strip():
                    pairs.append(f"{header}: {str(value).strip()}")
            if pairs:
                row_strings.append(" | ".join(pairs))

        sheet_texts.append("\n".join(row_strings))

    wb.close()
    return "\n\n".join(sheet_texts)


# ── Dispatcher ──────────────────────────────────────────────────────────────

EXTRACTORS = {
    ".pdf": extract_pdf,
    ".docx": extract_docx,
    ".txt": extract_txt,
    ".xlsx": extract_xlsx,
}


def extract_text(file_path: str) -> str:
    """
    Dispatch to the correct extractor based on file extension.
    
    Args:
        file_path: Path to the document file.
        
    Returns:
        Extracted text as a single string.
        
    Raises:
        ValueError: If the file format is not supported.
    """
    ext = Path(file_path).suffix.lower()
    if ext not in EXTRACTORS:
        supported = ", ".join(EXTRACTORS.keys())
        raise ValueError(f"Unsupported file format '{ext}'. Supported: {supported}")
    return EXTRACTORS[ext](file_path)
